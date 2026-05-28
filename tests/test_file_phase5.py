from io import BytesIO
from pathlib import Path
import socket
import sys

import pytest
from fastapi.testclient import TestClient

SERVER_ROOT = Path(__file__).resolve().parents[1] / "apps" / "server"
sys.path.insert(0, str(SERVER_ROOT))

from app.main import app  # noqa: E402


def require_local_port(port: int) -> None:
    try:
        with socket.create_connection(("127.0.0.1", port), timeout=1):
            return
    except OSError as exc:
        pytest.skip(f"localhost:{port} is not reachable: {exc}")


def test_upload_txt_parses_content_and_recommend_uses_file() -> None:
    require_local_port(5432)
    require_local_port(6379)

    with TestClient(app) as client:
        response = client.post(
            "/api/files/upload",
            files={"file": ("notes.txt", b"hello modelgate\nsecond line", "text/plain")},
        )

        assert response.status_code == 200
        file_record = response.json()["data"]
        assert file_record["status"] == "parsed"
        assert file_record["directUsable"] is True
        assert file_record["metadata"]["parsedText"].startswith("hello modelgate")
        assert file_record["previewUrl"] is None

        recommend_response = client.post(
            "/api/models/recommend",
            json={
                "taskType": "document_analysis",
                "inputTypes": ["text"],
                "fileIds": [file_record["id"]],
                "requiredOutput": "text",
            },
        )
        assert recommend_response.status_code == 200
        assert recommend_response.json()["data"]["availableModels"]


def test_rejects_mismatched_image_signature() -> None:
    require_local_port(5432)
    require_local_port(6379)

    with TestClient(app) as client:
        response = client.post(
            "/api/files/upload",
            files={"file": ("fake.png", b"not actually png", "image/png")},
        )

    assert response.status_code == 400
    assert response.json()["error"]["type"] == "FILE_SIGNATURE_MISMATCH"


def test_upload_docx_parses_paragraphs() -> None:
    require_local_port(5432)
    require_local_port(6379)

    from docx import Document

    buffer_path = Path("/private/tmp/modelgate-phase5-test.docx")
    document = Document()
    document.add_heading("Phase 5", level=1)
    document.add_paragraph("DOCX parser smoke test.")
    document.save(buffer_path)

    with TestClient(app) as client:
        response = client.post(
            "/api/files/upload",
            files={
                "file": (
                    "phase5.docx",
                    buffer_path.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    metadata = response.json()["data"]["metadata"]
    assert response.json()["data"]["status"] == "parsed"
    assert "DOCX parser smoke test." in metadata["parsedText"]
    assert metadata["paragraphCount"] >= 1


def test_upload_code_detects_language() -> None:
    require_local_port(5432)
    require_local_port(6379)

    with TestClient(app) as client:
        response = client.post(
            "/api/files/upload",
            files={"file": ("main.py", b"#!/usr/bin/env python\nprint('ok')\n", "text/x-python")},
        )

    assert response.status_code == 200
    metadata = response.json()["data"]["metadata"]
    assert response.json()["data"]["detectedType"] == "code"
    assert metadata["language"] == "python"
    assert metadata["lineCount"] == 2


def test_upload_markdown_extracts_headings_and_code_blocks() -> None:
    require_local_port(5432)
    require_local_port(6379)

    content = b"# Title\n\nSome notes.\n\n```ts\nconsole.log('ok')\n```\n"
    with TestClient(app) as client:
        response = client.post(
            "/api/files/upload",
            files={"file": ("notes.md", content, "text/markdown")},
        )

    assert response.status_code == 200
    metadata = response.json()["data"]["metadata"]
    assert metadata["language"] == "markdown"
    assert metadata["headings"] == ["Title"]
    assert metadata["codeBlockCount"] == 1


def test_upload_minimal_pdf_extracts_metadata() -> None:
    require_local_port(5432)
    require_local_port(6379)

    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    pdf_buffer = BytesIO()
    writer.write(pdf_buffer)

    with TestClient(app) as client:
        response = client.post(
            "/api/files/upload",
            files={"file": ("empty.pdf", pdf_buffer.getvalue(), "application/pdf")},
        )

    assert response.status_code == 200
    data = response.json()["data"]
    assert data["status"] == "parsed"
    assert data["metadata"]["pageCount"] == 1
