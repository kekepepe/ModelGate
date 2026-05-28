from __future__ import annotations

import csv
import io
import re
import zipfile
from dataclasses import dataclass
from html import unescape
from pathlib import Path
from uuid import uuid4

from app.core.config import settings

PARSE_VERSION = 1
MAX_PARSED_TEXT_CHARS = 200_000
MAX_CHUNK_TEXT_CHARS = 12_000
MAX_CSV_SAMPLE_ROWS = 20


@dataclass(frozen=True)
class ParseResult:
    status: str
    direct_usable: bool
    metadata: dict
    preview_path: str | None = None
    error_message: str | None = None


def parse_uploaded_file(
    *,
    stored_path: Path,
    original_name: str,
    mime_type: str,
    detected_type: str,
    safe_extension: str,
) -> ParseResult:
    base_metadata = {
        "originalName": original_name,
        "extension": safe_extension,
        "parser": "modelgate.local",
        "parseVersion": PARSE_VERSION,
    }

    try:
        if detected_type == "image":
            return _parse_image(stored_path=stored_path, metadata=base_metadata, mime_type=mime_type)
        if safe_extension == ".pdf":
            metadata = base_metadata | _parse_pdf(stored_path)
            return _parsed(metadata)
        if safe_extension == ".docx":
            metadata = base_metadata | _parse_docx(stored_path)
            return _parsed(metadata)
        if safe_extension == ".pptx":
            metadata = base_metadata | _parse_pptx(stored_path)
            return _parsed(metadata)
        if safe_extension == ".xlsx":
            metadata = base_metadata | _parse_xlsx(stored_path)
            return _parsed(metadata)
        if safe_extension == ".csv":
            metadata = base_metadata | _parse_csv(stored_path)
            return _parsed(metadata)
        if detected_type == "code" or safe_extension in {".txt", ".md"}:
            metadata = base_metadata | _parse_text_or_code(stored_path, safe_extension)
            return _parsed(metadata)
        if detected_type in {"video", "audio"}:
            media_type = "video" if detected_type == "video" else "audio"
            return _parsed(
                base_metadata
                | {
                    "mediaType": media_type,
                    "parsedText": "",
                    "chunks": [],
                    "note": "Basic media metadata only in Phase 5.",
                }
            )
        return _parsed(base_metadata | {"parsedText": "", "chunks": []})
    except Exception as exc:
        return ParseResult(
            status="failed",
            direct_usable=False,
            metadata=base_metadata,
            error_message=str(exc),
        )


def _parsed(metadata: dict) -> ParseResult:
    return ParseResult(status="parsed", direct_usable=True, metadata=metadata)


def _parse_image(*, stored_path: Path, metadata: dict, mime_type: str) -> ParseResult:
    if mime_type == "image/svg+xml" or stored_path.suffix.lower() == ".svg":
        text = _read_text(stored_path)
        return ParseResult(
            status="parsed",
            direct_usable=True,
            metadata=metadata
            | {
                "format": "SVG",
                "parsedText": _truncate(text, MAX_PARSED_TEXT_CHARS),
                "chunks": _line_chunks(text, chunk_type="svg"),
            },
            preview_path=None,
        )

    from PIL import Image

    with Image.open(stored_path) as image:
        width, height = image.size
        image_metadata = metadata | {
            "width": width,
            "height": height,
            "format": image.format,
            "mode": image.mode,
            "hasAlpha": image.mode in {"LA", "RGBA"} or "transparency" in image.info,
            "parsedText": "",
            "chunks": [],
        }

        preview_path = None
        if mime_type != "image/svg+xml":
            preview_root = _dated_root(Path(settings.previews_dir))
            preview_path_obj = preview_root / f"{uuid4().hex}.webp"
            image.thumbnail((1024, 1024))
            preview = Image.new("RGBA", image.size, (255, 255, 255, 0))
            converted = image.convert("RGBA")
            preview.paste(converted, (0, 0))
            preview.save(preview_path_obj, "WEBP", quality=82, exif=b"")
            preview_path = str(preview_path_obj)

    return ParseResult(
        status="parsed",
        direct_usable=True,
        metadata=image_metadata,
        preview_path=preview_path,
    )


def _parse_pdf(stored_path: Path) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(str(stored_path))
    if reader.is_encrypted:
        raise ValueError("Encrypted PDFs are not supported.")

    chunks = []
    pages_text = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages_text.append(text)
            chunks.append({"type": "page", "index": index, "text": _truncate(text, MAX_CHUNK_TEXT_CHARS)})

    parsed_text = _truncate("\n\n".join(pages_text), MAX_PARSED_TEXT_CHARS)
    return {"pageCount": len(reader.pages), "parsedText": parsed_text, "chunks": chunks}


def _parse_docx(stored_path: Path) -> dict:
    from docx import Document

    document = Document(str(stored_path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines = []
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_lines.append(f"Table {table_index} Row {row_index}: " + " | ".join(cells))

    full_text = "\n".join(paragraphs + table_lines)
    chunks = _line_chunks(full_text, chunk_type="docx")
    headings = [text for text in paragraphs if len(text) <= 120 and not text.endswith(".")][:20]
    return {
        "paragraphCount": len(paragraphs),
        "tableCount": len(document.tables),
        "headings": headings,
        "parsedText": _truncate(full_text, MAX_PARSED_TEXT_CHARS),
        "chunks": chunks,
    }


def _parse_pptx(stored_path: Path) -> dict:
    with zipfile.ZipFile(stored_path) as archive:
        slide_names = sorted(
            name
            for name in archive.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
        slide_texts = []
        chunks = []
        for index, slide_name in enumerate(slide_names, start=1):
            xml = archive.read(slide_name).decode("utf-8", errors="ignore")
            text = _extract_xml_text(xml)
            if text:
                slide_texts.append(text)
                chunks.append(
                    {"type": "slide", "index": index, "text": _truncate(text, MAX_CHUNK_TEXT_CHARS)}
                )
    return {
        "slideCount": len(slide_names),
        "parsedText": _truncate("\n\n".join(slide_texts), MAX_PARSED_TEXT_CHARS),
        "chunks": chunks,
    }


def _parse_xlsx(stored_path: Path) -> dict:
    from openpyxl import load_workbook

    workbook = load_workbook(filename=stored_path, read_only=True, data_only=True)
    sheet_summaries = []
    text_parts = []
    chunks = []
    for sheet in workbook.worksheets:
        rows = []
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                rows.append(values)
                text_parts.append(f"{sheet.title} R{row_index}: " + " | ".join(values))
            if len(rows) >= MAX_CSV_SAMPLE_ROWS:
                break
        sheet_summaries.append(
            {
                "name": sheet.title,
                "maxRow": sheet.max_row,
                "maxColumn": sheet.max_column,
                "sampleRows": rows,
            }
        )
        if rows:
            chunks.append(
                {
                    "type": "sheet",
                    "name": sheet.title,
                    "text": _truncate("\n".join(" | ".join(row) for row in rows), MAX_CHUNK_TEXT_CHARS),
                }
            )
    workbook.close()
    return {
        "sheetCount": len(sheet_summaries),
        "sheets": sheet_summaries,
        "parsedText": _truncate("\n".join(text_parts), MAX_PARSED_TEXT_CHARS),
        "chunks": chunks,
    }


def _parse_csv(stored_path: Path) -> dict:
    text = _read_text(stored_path)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except csv.Error:
        dialect = csv.excel

    reader = csv.reader(io.StringIO(text), dialect)
    rows = list(reader)
    headers = rows[0] if rows else []
    sample_rows = rows[1 : MAX_CSV_SAMPLE_ROWS + 1] if len(rows) > 1 else []
    column_count = max((len(row) for row in rows), default=0)
    parsed_lines = [" | ".join(row) for row in rows[: MAX_CSV_SAMPLE_ROWS + 1]]
    return {
        "headers": headers,
        "rowCount": max(len(rows) - 1, 0),
        "columnCount": column_count,
        "sampleRows": sample_rows,
        "parsedText": _truncate("\n".join(parsed_lines), MAX_PARSED_TEXT_CHARS),
        "chunks": _line_chunks("\n".join(parsed_lines), chunk_type="csv"),
    }


def _parse_text_or_code(stored_path: Path, extension: str) -> dict:
    text = _read_text(stored_path)
    language = _detect_language(stored_path.name, extension, text)
    headings = _extract_markdown_headings(text) if extension == ".md" else []
    code_blocks = len(re.findall(r"```", text)) // 2 if extension == ".md" else 0
    return {
        "language": language,
        "lineCount": len(text.splitlines()),
        "headings": headings,
        "codeBlockCount": code_blocks,
        "parsedText": _truncate(text, MAX_PARSED_TEXT_CHARS),
        "chunks": _line_chunks(text, chunk_type="code" if language else "text"),
    }


def _read_text(path: Path) -> str:
    content = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-16"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _extract_xml_text(xml: str) -> str:
    raw_parts = re.findall(r"<a:t[^>]*>(.*?)</a:t>", xml, flags=re.DOTALL)
    return "\n".join(unescape(re.sub(r"\s+", " ", part)).strip() for part in raw_parts if part.strip())


def _extract_markdown_headings(text: str) -> list[str]:
    headings = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            headings.append(stripped.lstrip("#").strip())
    return headings[:50]


def _line_chunks(text: str, *, chunk_type: str) -> list[dict]:
    chunks = []
    current: list[str] = []
    current_length = 0
    index = 1
    for line in text.splitlines():
        if current_length + len(line) > MAX_CHUNK_TEXT_CHARS and current:
            chunks.append({"type": chunk_type, "index": index, "text": "\n".join(current)})
            index += 1
            current = []
            current_length = 0
        current.append(line)
        current_length += len(line) + 1
    if current:
        chunks.append({"type": chunk_type, "index": index, "text": "\n".join(current)})
    return chunks[:50]


def _detect_language(filename: str, extension: str, text: str) -> str | None:
    special_names = {
        "dockerfile": "dockerfile",
        "makefile": "makefile",
    }
    lower_name = filename.lower()
    if lower_name in special_names:
        return special_names[lower_name]
    if text.startswith("#!"):
        first_line = text.splitlines()[0].lower()
        if "python" in first_line:
            return "python"
        if "node" in first_line:
            return "javascript"
        if "bash" in first_line or "sh" in first_line:
            return "shell"
    return {
        ".css": "css",
        ".go": "go",
        ".html": "html",
        ".java": "java",
        ".js": "javascript",
        ".json": "json",
        ".jsx": "javascriptreact",
        ".kt": "kotlin",
        ".md": "markdown",
        ".php": "php",
        ".py": "python",
        ".rs": "rust",
        ".sql": "sql",
        ".swift": "swift",
        ".toml": "toml",
        ".ts": "typescript",
        ".tsx": "typescriptreact",
        ".xml": "xml",
        ".yaml": "yaml",
        ".yml": "yaml",
    }.get(extension)


def _truncate(value: str, max_chars: int) -> str:
    if len(value) <= max_chars:
        return value
    return value[:max_chars]


def _dated_root(root: Path) -> Path:
    from datetime import datetime

    now = datetime.utcnow()
    path = root / f"{now:%Y}" / f"{now:%m}"
    path.mkdir(parents=True, exist_ok=True)
    return path
